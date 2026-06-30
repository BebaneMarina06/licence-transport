"""
Génère le document d'expression de besoin au format Word (.docx).
Usage : python scripts/generate_expression_besoin.py
"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

OUTPUT = Path(__file__).resolve().parent.parent / "docs" / "Expression_de_besoin_Licence_Transport_Gabon.docx"


def add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str, bold: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    return p


def add_bullet_list(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for row_data in rows:
        row = table.add_row().cells
        for i, cell in enumerate(row_data):
            row[i].text = cell
            for p in row[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()


def build_document() -> Document:
    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Page de garde
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("EXPRESSION DE BESOIN\n\n")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0, 85, 164)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Plateforme de digitalisation des licences de transport\n"
        "République Gabonaise\n\n"
    )
    run.font.size = Pt(14)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f"Version 1.0\n"
        f"Date : {date.today().strftime('%d/%m/%Y')}\n"
        "Maître d'ouvrage : Ministère des Transports\n"
        "Direction Générale des Transports Terrestres (DGTT)\n"
        "Porteur de projet : Ventis / Service Public Group"
    )
    run.font.size = Pt(11)

    doc.add_page_break()

    # Table des matières simplifiée
    add_heading(doc, "Table des matières", 1)
    toc_items = [
        "1. Contexte et problématique",
        "2. Objectifs du projet",
        "3. Périmètre et acteurs",
        "4. Besoins fonctionnels — Portail citoyen",
        "5. Besoins fonctionnels — Backoffice administration",
        "6. Types de licences et pièces justificatives",
        "7. Workflow de traitement des dossiers",
        "8. Exigences non fonctionnelles",
        "9. Architecture technique cible",
        "10. Sécurité et conformité",
        "11. Planning prévisionnel",
        "12. Risques et contraintes",
        "13. Critères d'acceptation",
    ]
    for item in toc_items:
        add_paragraph(doc, item)
    doc.add_page_break()

    # 1. Contexte
    add_heading(doc, "1. Contexte et problématique", 1)
    add_paragraph(
        doc,
        "Le Gabon engage une transformation numérique profonde de son administration publique, "
        "dans le cadre des initiatives « Gabon Digital », du Portail Gouvernemental des Services (PGS) "
        "et de la future plateforme « Gabon Connect ». Le permis de conduire digitalisé a été officiellement "
        "lancé en mars 2026 ; la digitalisation des cartes grises et des licences de transport constitue "
        "la prochaine étape structurante du secteur des transports terrestres.",
    )
    add_paragraph(
        doc,
        "À ce jour, l'obtention d'une licence de transport repose encore largement sur des procédures "
        "manuelles, des déplacements multiples et des circuits parallèles sources de fraude et de lenteurs. "
        "Les transporteurs professionnels (marchandises, voyageurs, activités mixtes) doivent se rendre "
        "physiquement aux services de la DGTT pour déposer leurs dossiers et régler les frais.",
    )
    add_paragraph(
        doc,
        "La présente expression de besoin définit les exigences pour la mise en place d'une plateforme "
        "numérique permettant aux citoyens et aux entreprises de souscrire, payer et suivre leurs licences "
        "de transport en ligne, tout en offrant aux agents administratifs un outil de gestion et d'instruction "
        "des dossiers.",
    )

    # 2. Objectifs
    add_heading(doc, "2. Objectifs du projet", 1)
    add_heading(doc, "2.1 Objectif général", 2)
    add_paragraph(
        doc,
        "Digitaliser l'ensemble du parcours de demande, de paiement et de délivrance des licences de "
        "transport terrestre au Gabon, en conformité avec le cadre réglementaire national et les tarifs "
        "officiels publiés par l'administration.",
    )
    add_heading(doc, "2.2 Objectifs spécifiques", 2)
    add_bullet_list(
        doc,
        [
            "Permettre aux usagers de créer un compte et d'effectuer leurs démarches à distance ;",
            "Réduire les délais de traitement et les déplacements physiques ;",
            "Sécuriser les documents et lutter contre la fraude documentaire ;",
            "Assurer la traçabilité complète des dossiers et des paiements ;",
            "Offrir aux agents un backoffice d'instruction, de validation et de délivrance ;",
            "Préparer l'interopérabilité avec l'écosystème national (ANINF, Gabon Connect, permis digitalisé).",
        ],
    )

    # 3. Périmètre
    add_heading(doc, "3. Périmètre et acteurs", 1)
    add_heading(doc, "3.1 Périmètre inclus", 2)
    add_bullet_list(
        doc,
        [
            "Portail citoyen / transporteur (web responsive) ;",
            "Backoffice administration DGTT (web) ;",
            "API centrale de gestion des dossiers ;",
            "Gestion des 5 types de licences officiels ;",
            "Upload et consultation des pièces justificatives ;",
            "Workflow d'instruction et de délivrance ;",
            "Tableau de bord et statistiques.",
        ],
    )
    add_heading(doc, "3.2 Périmètre hors phase initiale", 2)
    add_bullet_list(
        doc,
        [
            "Paiement Mobile Money (phase 2) ;",
            "Génération de licence PDF avec QR code / NFC (phase 2) ;",
            "Notifications SMS / email automatisées (phase 2) ;",
            "SSO Gabon Connect (phase 3) ;",
            "Interconnexion avec le fichier permis de conduire et carte grise (phase 3).",
        ],
    )
    add_heading(doc, "3.3 Acteurs", 2)
    add_table(
        doc,
        ["Acteur", "Rôle"],
        [
            ["Citoyen / transporteur", "Dépose sa demande, joint les pièces, paie, suit son dossier"],
            ["Agent DGTT", "Instruit les dossiers, demande des compléments, valide ou rejette"],
            ["Superviseur", "Supervise l'activité, gère les files d'attente"],
            ["Administrateur", "Gère les utilisateurs, les paramètres, les types de licences"],
            ["Auditeur", "Consulte les journaux d'audit et les statistiques"],
        ],
    )

    # 4. Portail citoyen
    add_heading(doc, "4. Besoins fonctionnels — Portail citoyen", 1)
    add_table(
        doc,
        ["ID", "Besoin", "Priorité", "Statut"],
        [
            ["BC-01", "Page d'accueil présentant les types de licences et tarifs officiels", "Haute", "Réalisé"],
            ["BC-02", "Inscription et connexion sécurisées (email / mot de passe)", "Haute", "Réalisé"],
            ["BC-03", "Espace personnel listant les dossiers de l'usager", "Haute", "Réalisé"],
            ["BC-04", "Création d'une demande de licence (formulaire guidé)", "Haute", "Réalisé"],
            ["BC-05", "Upload des 3 justificatifs obligatoires (carte grise, visite technique, assurance)", "Haute", "Réalisé"],
            ["BC-06", "Soumission du dossier après complétude des pièces", "Haute", "Réalisé"],
            ["BC-07", "Suivi en temps réel avec historique des statuts", "Haute", "Réalisé"],
            ["BC-08", "Paiement en ligne des frais de licence", "Haute", "À faire"],
            ["BC-09", "Téléchargement de la licence délivrée (PDF + QR code)", "Haute", "À faire"],
            ["BC-10", "Notifications par SMS / email aux changements de statut", "Moyenne", "À faire"],
            ["BC-11", "Renouvellement de licence existante", "Moyenne", "À faire"],
            ["BC-12", "Interface mobile-first et multilingue (FR/EN)", "Moyenne", "À faire"],
        ],
    )

    # 5. Backoffice
    add_heading(doc, "5. Besoins fonctionnels — Backoffice administration", 1)
    add_table(
        doc,
        ["ID", "Besoin", "Priorité", "Statut"],
        [
            ["BA-01", "Connexion réservée au personnel administratif (RBAC)", "Haute", "Réalisé"],
            ["BA-02", "Tableau de bord avec indicateurs (dossiers en attente, délivrés, recettes)", "Haute", "Réalisé"],
            ["BA-03", "File d'instruction filtrable par statut et type de licence", "Haute", "Réalisé"],
            ["BA-04", "Consultation du dossier et des pièces justificatives", "Haute", "Réalisé"],
            ["BA-05", "Workflow : prise en charge, complément, approbation, rejet", "Haute", "Réalisé"],
            ["BA-06", "Validation du paiement et délivrance de la licence", "Haute", "Réalisé"],
            ["BA-07", "Journal d'audit des actions agents", "Haute", "Réalisé"],
            ["BA-08", "Gestion des utilisateurs et des rôles", "Moyenne", "À faire"],
            ["BA-09", "Export des statistiques (Excel / PDF)", "Moyenne", "À faire"],
            ["BA-10", "Génération et impression de la licence sécurisée", "Haute", "À faire"],
        ],
    )

    # 6. Licences
    add_heading(doc, "6. Types de licences et pièces justificatives", 1)
    add_paragraph(
        doc,
        "Source officielle : https://infrastructures.gouv.ga/18-transport/20-documents-administratifs/329-licence-de-transport/",
    )
    add_table(
        doc,
        ["Type de licence", "Tarif (F CFA)", "Justificatifs requis"],
        [
            ["Autorisation Spéciale de Transport", "50 000", "Carte grise, Visite technique, Assurance"],
            ["Licence Mixte", "200 000", "Carte grise, Visite technique, Assurance"],
            ["Licence transports voyageurs", "150 000 (*)", "Carte grise, Visite technique, Assurance"],
            ["Licence transports marchandises", "300 000 (**)", "Carte grise, Visite technique, Assurance"],
            ["Licence Exceptionnelle", "400 000", "Carte grise, Visite technique, Assurance"],
        ],
    )
    add_paragraph(doc, "(*) Tarif pour véhicule de moins de 18 places.", bold=False)
    add_paragraph(doc, "(**) Tarif variable selon le tonnage du véhicule.", bold=False)
    add_paragraph(
        doc,
        "Formats de fichiers acceptés pour l'upload : PDF, JPG, PNG. Taille maximale : 5 Mo par document.",
    )

    # 7. Workflow
    add_heading(doc, "7. Workflow de traitement des dossiers", 1)
    add_paragraph(doc, "Le cycle de vie d'un dossier suit les étapes suivantes :")
    add_bullet_list(
        doc,
        [
            "Brouillon — le demandeur complète son formulaire et joint les pièces ;",
            "Soumis — le dossier est transmis à l'administration ;",
            "En instruction — un agent prend en charge le dossier ;",
            "Complément demandé — pièces ou informations manquantes ;",
            "Approuvé — le dossier est validé sur le fond ;",
            "En attente de paiement — le demandeur doit régler les frais ;",
            "Payé — paiement confirmé ;",
            "Délivré — licence générée et mise à disposition ;",
            "Rejeté / Annulé — dossier refusé avec motif.",
        ],
    )

    # 8. ENF
    add_heading(doc, "8. Exigences non fonctionnelles", 1)
    add_bullet_list(
        doc,
        [
            "Disponibilité : 99,5 % en production ;",
            "Temps de réponse API < 2 secondes pour 95 % des requêtes ;",
            "Interface responsive (smartphone, tablette, desktop) ;",
            "Accessibilité et simplicité d'usage (logique Service Public) ;",
            "Hébergement souverain compatible infrastructure ANINF ;",
            "Sauvegarde quotidienne des données et des fichiers ;",
            "Scalabilité pour extension à d'autres titres de transport.",
        ],
    )

    # 9. Architecture
    add_heading(doc, "9. Architecture technique cible", 1)
    add_table(
        doc,
        ["Composant", "Technologie", "Description"],
        [
            ["Portail citoyen", "React + TypeScript + Vite", "Interface usager"],
            ["Backoffice", "React + TypeScript + Vite", "Interface administration"],
            ["API", "FastAPI (Python)", "Logique métier, auth, workflow"],
            ["Base de données", "PostgreSQL 16", "Données relationnelles"],
            ["Cache", "Redis", "Sessions et files d'attente"],
            ["Stockage fichiers", "MinIO / S3", "Pièces justificatives et licences PDF"],
            ["Conteneurisation", "Docker", "Déploiement reproductible"],
        ],
    )

    # 10. Sécurité
    add_heading(doc, "10. Sécurité et conformité", 1)
    add_bullet_list(
        doc,
        [
            "Authentification JWT avec tokens d'accès et de rafraîchissement ;",
            "Contrôle d'accès par rôles (RBAC) ;",
            "Chiffrement HTTPS obligatoire ;",
            "Journal d'audit immuable des actions sensibles ;",
            "Protection des données personnelles (conformité cadre e-gouvernement gabonais) ;",
            "Documents officiels avec QR code signé pour vérification terrain ;",
            "Authentification forte (2FA) pour les agents administratifs.",
        ],
    )

    # 11. Planning
    add_heading(doc, "11. Planning prévisionnel", 1)
    add_table(
        doc,
        ["Phase", "Durée estimée", "Livrables"],
        [
            ["Phase 0 — Fondations", "2 semaines", "Architecture, auth, modèles, CI/CD"],
            ["Phase 1 — MVP", "4 semaines", "Dépôt dossier, upload pièces, instruction backoffice"],
            ["Phase 2 — Paiement & délivrance", "3 semaines", "Mobile Money, PDF licence QR code"],
            ["Phase 3 — Notifications & reporting", "2 semaines", "SMS, email, exports statistiques"],
            ["Phase 4 — Interopérabilité", "Continu", "Gabon Connect, ANINF, permis / carte grise"],
        ],
    )

    # 12. Risques
    add_heading(doc, "12. Risques et contraintes", 1)
    add_table(
        doc,
        ["Risque", "Impact", "Mitigation"],
        [
            ["Connectivité limitée hors Libreville", "Élevé", "Interface légère, mode dégradé, France Services"],
            ["Adoption Mobile Money hétérogène", "Moyen", "Module paiement abstrait multi-opérateurs"],
            ["Résistance au changement administratif", "Moyen", "Formation agents, conduite du changement"],
            ["Fraude documentaire", "Élevé", "QR code, vérification en ligne, audit"],
            ["Charge pic de demandes", "Moyen", "Scalabilité cloud, files d'attente"],
        ],
    )

    # 13. Critères
    add_heading(doc, "13. Critères d'acceptation", 1)
    add_bullet_list(
        doc,
        [
            "Un citoyen peut s'inscrire, créer un dossier, joindre les 3 pièces et le soumettre sans erreur ;",
            "Un agent peut instruire, approuver ou rejeter un dossier avec traçabilité ;",
            "Les tarifs affichés correspondent aux tarifs officiels du site gouvernemental ;",
            "Les documents uploadés sont consultables par les agents autorisés ;",
            "Le tableau de bord reflète les statistiques en temps réel ;",
            "Les tests de sécurité (auth, RBAC, audit) sont validés avant mise en production.",
        ],
    )

    doc.add_paragraph()
    add_paragraph(doc, "— Fin du document —", bold=True)

    return doc


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Document généré : {OUTPUT}")


if __name__ == "__main__":
    main()
